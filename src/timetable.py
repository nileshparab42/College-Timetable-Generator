import random
import docx
import os 

# FOR TAKING INPUTS OF TEACHERS INFO
print("===============================================")
noofteachers = int(input("ENTER NUMBER OF THEACHERS : "))
print("===============================================")
os.system('cls')
teachers=[[] for _ in range(0,int(noofteachers))]
for i in range(0,int(noofteachers)):
    print("===============================================")
    print("Enter info of teacher "+str(i+1)+" :")
    print("===============================================")
    teachers[i].append(input("Name : "))
    teachers[i].append(input("Time Slot : "))
    print("-----------------------------------------------")
    teachers[i].append(input("subject(SE) : "))
    teachers[i].append(int(input("Th(SE) : ")))
    teachers[i].append(int(input("PT(SE) : ")))
    teachers[i].append(input("Enter lab room no(SE) : "))
    print("-----------------------------------------------")
    teachers[i].append(input("subject(TE) : "))
    teachers[i].append(int(input("Th(TE) : ")))
    teachers[i].append(int(input("PT(TE) : ")))
    teachers[i].append(input("Enter lab room no(TE) : "))
    print("-----------------------------------------------")
    teachers[i].append(input("subject(BE) : "))
    teachers[i].append(int(input("Th(BE) : ")))
    teachers[i].append(int(input("PT(BE) : ")))
    teachers[i].append(input("Enter lab room no(BE) : "))
    print("===============================================")
    os.system('cls')

# SAMPLE OF TEACHERS FOR TESTING
# noofteachers=16
# teachers=[["CG","C","OS",3,2,"704","",0,0,"","",0,0,""],["NM","A","MP",3,2,"701","",0,0,"","",0,0,""],["AG","C","DBMS",3,2,"LL","",0,0,"","",0,0,""],["DR","B","AOA",3,2,"702","",0,0,"","",0,0,""],["AS","C","PYTHON",2,2,"501","",0,0,"","",0,0,""],["AJ","D","EM4",3,0,"","",0,0,"","",0,0,""],["RJ","A","",0,0,"","SPCC",3,2,"704","",0,0,""],["NV","C","",0,0,"","AI",3,2,"501","",0,0,""],["UA","C","",0,0,"","CNSS",3,2,"108","",0,0,""],["RS","C","",0,0,"","MC",3,2,"704","",0,0,""],["DNR","A","",0,0,"","IOT",3,0,"","",0,0,""],["MNR","C","",0,0,"","CC",0,2,"","",0,0,""],["SS","A","",0,0,"","",0,0,"","HMI",4,2,"702"],["VJ","C","",0,0,"","",0,0,"","EM",3,0,""],["TRP","A","",0,0,"","",0,0,"","AWN",4,2,"501"],["SS","B","",0,0,"","",0,0,"","DC",4,2,""]]

# STRUCTURE OF TT
se_tt=[[["LAB",""],["LAB",""],[],[],[],[]],[[],[],["LAB",""],["LAB",""],[],[]],[[],[],[],[],[],[]],[[],[],[],[],[],[]],[["LAB",""],["LAB",""],[],[],[],[]],[["LAB",""],["LAB",""],["LAB",""],["LAB",""],[],[]]]

te_tt=[[[],[],[],[],["LAB",""],["LAB",""]],[[],[],[],[],["LAB",""],["LAB",""]],[[],[],["LAB",""],["LAB",""],["LAB",""],["LAB",""]],[[],[],[],[],["LAB",""],["LAB",""]],[[],[],["LAB",""],["LAB",""],["LAB",""],["LAB",""]],[[],[],[],[],[],[]]]

be_tt=[[[],[],["LAB",""],["LAB",""],[],[]],[["LAB",""],["LAB",""],[],[],[],[]],[["LAB",""],["LAB",""],[],[],["LAB",""],["LAB",""]],[[],[],["LAB",""],["LAB",""],[],[]],[[],[],[],[],[],[]],[[],[],[],[],[],[]]]

# SHUFFLE TEACHER LIST
random.shuffle(teachers)

# ARRANGING SE TIMETABLE
for i in range(0,6):
    for j in range(0,6):
        if len(se_tt[j][i])==0:
            if i==0:
                random.shuffle(teachers)
                for z in range(0,int(noofteachers)):
                    if teachers[z][1]=="A" :
                        if teachers[z][3]!=0:
                            teachers[z][3]=teachers[z][3]-1
                            se_tt[j][i].append(teachers[z][2])
                            se_tt[j][i].append(teachers[z][0])
                            break
            elif i==5:
                random.shuffle(teachers)
                for z in range(0,int(noofteachers)):
                    if teachers[z][1]=="B" :
                        if teachers[z][3]!=0:
                            teachers[z][3]=teachers[z][3]-1
                            se_tt[j][i].append(teachers[z][2])
                            se_tt[j][i].append(teachers[z][0])
                            break
            else:
                random.shuffle(teachers)
                for z in range(0,int(noofteachers)):
                    if teachers[z][3]!=0:
                        teachers[z][3]=teachers[z][3]-1
                        se_tt[j][i].append(teachers[z][2])
                        se_tt[j][i].append(teachers[z][0])
                        break

# ARRANGING TE TIMETABLE
for i in range(0,6):
    for j in range(0,5):
        if len(te_tt[j][i])==0:
            if i==0:
                random.shuffle(teachers)
                for z in range(0,int(noofteachers)):
                    if teachers[z][1]=="A" :
                        if teachers[z][7]!=0:
                            teachers[z][7]=teachers[z][7]-1
                            te_tt[j][i].append(teachers[z][6])
                            te_tt[j][i].append(teachers[z][0])
                            break
            elif i==5:
                random.shuffle(teachers)
                for z in range(0,int(noofteachers)):
                    if teachers[z][1]=="B" :
                        if teachers[z][7]!=0:
                            teachers[z][7]=teachers[z][7]-1
                            te_tt[j][i].append(teachers[z][6])
                            te_tt[j][i].append(teachers[z][0])
                            break
            else:
                random.shuffle(teachers)
                for z in range(0,int(noofteachers)):
                    if teachers[z][7]!=0:
                        teachers[z][7]=teachers[z][7]-1
                        te_tt[j][i].append(teachers[z][6])
                        te_tt[j][i].append(teachers[z][0])
                        break
                            

# ARRANGING BE TIMETABLE
for i in range(0,6):
    for j in range(0,5):
        if len(be_tt[j][i])==0:
            if i==0:
                random.shuffle(teachers)
                for z in range(0,int(noofteachers)):
                    if teachers[z][1]=="A" :
                        if teachers[z][11]!=0:
                            teachers[z][11]=teachers[z][11]-1
                            be_tt[j][i].append(teachers[z][10])
                            be_tt[j][i].append(teachers[z][0])
                            break
            elif i==5:
                random.shuffle(teachers)
                for z in range(0,int(noofteachers)):
                    if teachers[z][1]=="B" :
                        if teachers[z][11]!=0:
                            teachers[z][11]=teachers[z][11]-1
                            be_tt[j][i].append(teachers[z][10])
                            be_tt[j][i].append(teachers[z][0])
                            break
            else:
                random.shuffle(teachers)
                for z in range(0,int(noofteachers)):
                    if teachers[z][11]!=0:
                        teachers[z][11]=teachers[z][11]-1
                        be_tt[j][i].append(teachers[z][10])
                        be_tt[j][i].append(teachers[z][0])
                        break

# PRINTING ARRANGED TIMETABLES           
print("SE TIMETABLE :")
print(se_tt)
print("\n\n")
print("TE TIMETABLE :")
print(te_tt)
print("\n\n")
print("BE TIMETABLE :")
print(be_tt)

# MAPPING DATA TO WORD FILE
days=["MON","TUE","WED","THU","FRI","SAT"]

# MAPPING SE TIME TABLE
doc1=docx.Document()
doc1.add_heading("Second Year Timetable",0) 

se_timetable = doc1.add_table(rows=1,cols=7)
se_timetable.style="Table Grid"
hdr_cells=se_timetable.rows[0].cells

hdr_cells[1].text="9:45-10:45"
hdr_cells[2].text="10:45-11:45"
hdr_cells[3].text="11:45-12:45"
hdr_cells[4].text="12:45-1:45"
hdr_cells[5].text="2:15-3:15"
hdr_cells[6].text="3:15-4:45"

day=0

for lec1,lec2,lec3,lec4,lec5,lec6 in se_tt:
    row_cells=se_timetable.add_row().cells
    row_cells[0].text=days[day]
    day=day+1
    for i in lec1:
        row_cells[1].text=lec1[0]+"("+lec1[1]+")"
    for i in lec2:
        row_cells[2].text=lec2[0]+"("+lec2[1]+")"
    for i in lec3:
        row_cells[3].text=lec3[0]+"("+lec3[1]+")"
    for i in lec4:
        row_cells[4].text=lec4[0]+"("+lec4[1]+")"
    for i in lec5:
        row_cells[5].text=lec5[0]+"("+lec5[1]+")"
    for i in lec6:
        row_cells[6].text=lec6[0]+"("+lec6[1]+")"

doc1.save("se_timetable.docx")
os.system("start se_timetable.docx")

# MAPPING TE TIME TABLE
doc2=docx.Document()
doc2.add_heading("Third Year Timetable",0) 

te_timetable = doc2.add_table(rows=1,cols=7)
te_timetable.style="Table Grid"
hdr_cells=te_timetable.rows[0].cells

hdr_cells[1].text="9:45-10:45"
hdr_cells[2].text="10:45-11:45"
hdr_cells[3].text="11:45-12:45"
hdr_cells[4].text="12:45-1:45"
hdr_cells[5].text="2:15-3:15"
hdr_cells[6].text="3:15-4:45"

day=0

for lec1,lec2,lec3,lec4,lec5,lec6 in te_tt:
    row_cells=te_timetable.add_row().cells
    row_cells[0].text=days[day]
    day=day+1
    for i in lec1:
        row_cells[1].text=lec1[0]+"("+lec1[1]+")"
    for i in lec2:
        row_cells[2].text=lec2[0]+"("+lec2[1]+")"
    for i in lec3:
        row_cells[3].text=lec3[0]+"("+lec3[1]+")"
    for i in lec4:
        row_cells[4].text=lec4[0]+"("+lec4[1]+")"
    for i in lec5:
        row_cells[5].text=lec5[0]+"("+lec5[1]+")"
    for i in lec6:
        row_cells[6].text=lec6[0]+"("+lec6[1]+")"

doc2.save("te_timetable.docx")
os.system("start te_timetable.docx")

# MAPPING BE TIME TABLE
doc3=docx.Document()
doc3.add_heading("Final Year Timetable",0) 

be_timetable = doc3.add_table(rows=1,cols=7)
be_timetable.style="Table Grid"
hdr_cells=be_timetable.rows[0].cells

hdr_cells[1].text="9:45-10:45"
hdr_cells[2].text="10:45-11:45"
hdr_cells[3].text="11:45-12:45"
hdr_cells[4].text="12:45-1:45"
hdr_cells[5].text="2:15-3:15"
hdr_cells[6].text="3:15-4:45"

day=0

for lec1,lec2,lec3,lec4,lec5,lec6 in be_tt:
    row_cells=be_timetable.add_row().cells
    row_cells[0].text=days[day]
    day=day+1
    for i in lec1:
        row_cells[1].text=lec1[0]+"("+lec1[1]+")"
    for i in lec2:
        row_cells[2].text=lec2[0]+"("+lec2[1]+")"
    for i in lec3:
        row_cells[3].text=lec3[0]+"("+lec3[1]+")"
    for i in lec4:
        row_cells[4].text=lec4[0]+"("+lec4[1]+")"
    for i in lec5:
        row_cells[5].text=lec5[0]+"("+lec5[1]+")"
    for i in lec6:
        row_cells[6].text=lec6[0]+"("+lec6[1]+")"

doc3.save("be_timetable.docx")
os.system("start be_timetable.docx")

